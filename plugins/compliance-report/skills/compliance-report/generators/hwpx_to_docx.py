# -*- coding: utf-8 -*-
"""준법사항체크리스트 HWPX → DOCX(준법감시보고서) 변환 모듈.

hwpx_generator 가 생성한 양식 HWPX 를 한컴오피스 COM 자동화로 열어
MS Word(.docx) 형식으로 저장한다. 표·서식(적색 강조 등)이 그대로 보존된다.

전제:
    - Windows + 한컴오피스(HWPFrame.HwpObject) 설치
    - docx 저장 필터 문자열은 "OOXML"

알려진 한컴 COM 변환 한계 보정:
    - 별첨2 「여신/외환/수신」 Y/N 표는 헤더 표 셀 안에 중첩(nested table)되어 있어
      docx 내보내기 시 누락된다. _repair_attach2_table() 가 변환 후 docx 에
      해당 표를 동일 구조로 재삽입한다.
"""
import os
import sys


def convert_hwpx_to_docx(hwpx_path: str, docx_path: str) -> bool:
    """HWPX 를 DOCX 로 변환. 성공 시 True.

    한컴오피스가 없거나 변환 실패 시 False 반환(예외는 메시지 출력 후 흡수).
    """
    hwpx_path = os.path.abspath(hwpx_path)
    docx_path = os.path.abspath(docx_path)
    if not os.path.exists(hwpx_path):
        print(f"[ERROR] HWPX 파일이 없습니다: {hwpx_path}")
        return False

    # pywin32 사용자 설치 경로 보강 (전역 지침과 동일)
    user_site = r"C:\Users\anton\AppData\Roaming\Python\Python314\site-packages"
    if os.path.isdir(user_site) and user_site not in sys.path:
        sys.path.insert(0, user_site)

    try:
        import win32com.client as win32
    except ImportError:
        print("[ERROR] pywin32(win32com) 가 필요합니다. docx 변환을 건너뜁니다.")
        return False

    os.makedirs(os.path.dirname(docx_path) or ".", exist_ok=True)

    hwp = None
    saved = False
    try:
        hwp = win32.Dispatch("HWPFrame.HwpObject")
        hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        hwp.Open(hwpx_path, "HWPX", "forceopen:true")
        ok = hwp.SaveAs(docx_path, "OOXML", "")
        saved = bool(ok) and os.path.exists(docx_path)
        if not saved:
            print(f"[ERROR] DOCX 저장 실패: {docx_path}")
            return False
    except Exception as e:  # noqa: BLE001
        print(f"[ERROR] HWPX→DOCX 변환 중 오류: {e}")
        return False
    finally:
        if hwp is not None:
            try:
                hwp.Quit()
            except Exception:  # noqa: BLE001
                pass

    # 한컴이 파일을 닫은 뒤(잠금 해제) 별첨2 Y/N 표 재삽입
    if saved:
        _repair_attach2_table(docx_path)
    return saved


def hancom_clean_hwpx(hwpx_path: str) -> bool:
    """ZIP 기반으로 생성한 HWPX 를 한컴으로 열어(복구 로드) HWPX 로 재저장한다.

    파이썬 ZIP 치환본은 한컴이 '복구' 모드로만 열리는 경우가 있어(파일 자체는
    유효 XML 이지만 한컴의 내부 검증을 통과하지 못함) 사용자가 더블클릭하면
    복구 대화상자가 뜬다. 한컴이 한 번 로드 후 네이티브로 재저장하면 그 산출물은
    복구 없이 정상적으로 열린다.

    한컴이 없거나 실패하면 원본을 그대로 두고 False 를 반환한다(치명적이지 않음).
    """
    hwpx_path = os.path.abspath(hwpx_path)
    if not os.path.exists(hwpx_path):
        return False

    user_site = r"C:\Users\anton\AppData\Roaming\Python\Python314\site-packages"
    if os.path.isdir(user_site) and user_site not in sys.path:
        sys.path.insert(0, user_site)
    try:
        import win32com.client as win32
    except ImportError:
        print("[clean] pywin32 미설치 — HWPX 정리(round-trip)를 건너뜁니다.")
        return False

    tmp = hwpx_path + ".clean.hwpx"
    hwp = None
    ok = False
    try:
        hwp = win32.Dispatch("HWPFrame.HwpObject")
        hwp.RegisterModule("FilePathCheckDLL", "FilePathCheckerModule")
        # 모든 대화상자(복구·저장 등) 자동 처리 → 사용자 팝업 없음
        try:
            hwp.SetMessageBoxMode(0x00020000)
        except Exception:  # noqa: BLE001
            pass
        hwp.Open(hwpx_path, "HWPX", "forceopen:true")   # 복구 로드(반환값 무시)
        ok = bool(hwp.SaveAs(tmp, "HWPX", "")) and os.path.exists(tmp)
        try:
            hwp.Run("FileClose")
        except Exception:  # noqa: BLE001
            pass
    except Exception as e:  # noqa: BLE001
        print(f"[clean] HWPX round-trip 중 오류: {e}")
        ok = False
    finally:
        if hwp is not None:
            try:
                hwp.Quit()
            except Exception:  # noqa: BLE001
                pass

    if ok:
        try:
            os.replace(tmp, hwpx_path)
        except OSError:
            ok = False
    if not ok and os.path.exists(tmp):
        try:
            os.remove(tmp)
        except OSError:
            pass
    return ok


def _apply_table_borders(table, qn):
    """표 전체에 단선 테두리(0.5pt)를 OXML 로 직접 적용 (명명 스타일 불요)."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def _repair_attach2_table(docx_path: str) -> bool:
    """별첨2 「여신/외환/수신」 Y/N 표가 변환 중 누락된 경우 재삽입.

    표 구조(3×4, 라벨 고정 / 값 셀 공란):
        구 분 | 여신 | 외환 | 수신*
        (Y/N) |      |      |
        중소기업은행 거래지점 |      | 증빙서류 |
    이미 존재하면(멱등) 아무 것도 하지 않는다.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        print("[주의] python-docx 미설치 — 별첨2 표 보정을 건너뜁니다.")
        return False

    doc = Document(docx_path)

    # 이미 여신 표가 있으면 보정 불필요
    full = "|".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    if "여신" in full and "외환" in full:
        return True

    # 헤더 표(■ 중소기업은행 ... 인정 여부) 탐색
    anchor = None
    for t in doc.tables:
        tx = "".join(c.text for r in t.rows for c in r.cells)
        if "중소기업은행 신규 및 기존거래 기업 인정 여부" in tx:
            anchor = t
            break
    if anchor is None:
        print("[주의] 별첨2 헤더 표를 찾지 못해 표 보정을 건너뜁니다.")
        return False

    rows = [
        ["구 분", "여신", "외환", "수신*"],
        ["(Y/N)", "", "", ""],
        ["중소기업은행 거래지점", "", "증빙서류", ""],
    ]
    new_tbl = doc.add_table(rows=len(rows), cols=4)
    _apply_table_borders(new_tbl, qn)
    for ri, rowvals in enumerate(rows):
        for ci, val in enumerate(rowvals):
            cell = new_tbl.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.name = "맑은 고딕"
            run.font.size = Pt(9)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.append(rfonts)
            for a in ("w:eastAsia", "w:ascii", "w:hAnsi"):
                rfonts.set(qn(a), "맑은 고딕")
            # 라벨 셀(0,2번 열의 텍스트 라벨)은 볼드
            if (ci == 0) or (ri == 2 and ci == 2) or ri == 0:
                run.font.bold = True

    # 헤더 표 바로 뒤로 이동
    anchor._tbl.addnext(new_tbl._tbl)
    doc.save(docx_path)
    print("[OK] 별첨2 여신/외환/수신 표 재삽입 완료")
    return True


if __name__ == "__main__":
    if len(sys.argv) < 3:
        sys.exit("사용법: python hwpx_to_docx.py <input.hwpx> <output.docx>")
    success = convert_hwpx_to_docx(sys.argv[1], sys.argv[2])
    print("완료" if success else "실패")
    sys.exit(0 if success else 1)
